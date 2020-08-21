#! python3
from docx import Document
from copy import deepcopy

                            #######################
                            # Create Dictionaries #
                            #######################
titleDict = {
    'THE CARBONARO EFFECT': 'CARBONARO EFFECT',
    'PARAMOUNT PLATINUM HD FILM PACKAGE': 'PARAMOUNT PLATINUM HD',
    'MIRAMAX ONE FILM PACKAGE': 'MIRAMAX ONE',
    'MIRAMAX FILM PACKAGE': 'MIRAMAX ONE',
    '2020 HOLIDAY FILM PACKAGE': 'HOLIDAY FIVE',
    '\“WHACKED OUT SPORTS\” DEAL MEMO' : 'WOS',
    '“WHACKED OUT SPORTS” DEAL MEMO' : 'WOS'
}

monDict = {
    'January': '01',
    'February': '02',
    'Februyary': '02',
    'March': '03',
    'April': '04',
    'May': '05',
    'June': '06',
    'July': '07',
    'August': '08',
    'September': '09',
    'October': '10',
    'November': '11',
    'December': '12',
}

markDict = {
    'Albany, NY': 'Albany-Schenectady',
    'Albany': 'Albany-Schenectady',
    'Albuquerque-Santa Fe': 'Albuquerque',
    'Burlington-Plattsburgh': 'Burlington - Plattsburgh',
    'Bulrington-Plattsburgh': 'Burlington - Plattsburgh',
    'Cedar Rapids – Waterloo': 'Cedar Rapids',
    'Champaign-Springfield': 'Champaign-Springfld',
    'Chico-Redding': 'Chico - Redding',
    'Columbia-Jefferson City': 'Columbia/Jeff City',
    'Davenport-Moline': 'Davenport',
    'Fort Myers': 'Ft. Myers - Naples',
    'Ft. Meyers': 'Ft. Myers - Naples',
    'Florence': 'Florence - Myrtle Beach',
    'Greensboro': 'Greensboro - H. Point',
    'Greenville-Spartanburg': 'GSA',
    'Greenville-New Bern': 'Greenville - N. Bern',
    'Harlingen-McAllen': 'Harlingen',
    'Hartford-New Haven': 'Hartford & New Haven',
    'Hattiesburg': 'Hattiesburg - Laurel',
    'Huntsville': 'Huntsville - Decatur',
    'Johnstown-Altoona': 'Johnstown - Altoona',
    'Lincoln-Hastings': 'Lincoln',
    'Medford': 'Medford - Klamath Falls',
    'Monroe': 'Monroe -  El Dorado',
    'Monterey': 'Monterey-Salinas',
    'Morgan City': 'Morgan City, LA',
    'Myrtle Beach': 'Florence - Myrtle Beach',
    'Norfolk': 'Norfolk - Portsmouth',
    'Orlando': 'Orlando - Daytona',
    'Portland': 'Portland, OR',
    'Providence': 'Providence-Nw Bedford',
    'Raleigh': 'Raleigh - Durham',
    'Richmond': 'Richmond-Petersburg',
    'Roanoke-Lynchburg': 'Roanoke - Lynchburg',
    'Roanoke': 'Roanoke - Lynchburg',
    'Rochester-Mason City-Austin': 'Rochester/      Mason CTY',
    'Waco': 'Waco-Temple',
    'ichita Falls': 'Wichita Falls',
    'Wilkes Barre-Scranton': 'Wilkes Barre - Scranton',
}

                                 #############
                                 # Functions #
                                 #############

def getInfo(filename):
    """Takes in a document and returns completed cell dictionary"""
    paras = deepcopy(Document(filename).paragraphs)
    cell = {'TITLE': '', 'DATE': '', 'STATION': '', 'MARKET': ''}

    for line in paras:
        if (line.text == '' or line.text == ' '):
            continue
        elif (cell.get('TITLE') == '' and line.text != ''):
            cell = assoc(cell, 'TITLE', cleanTitle(line.text))
        elif (line.text.find('DATED') > -1):
            cell = assoc(cell, 'DATE', cleanDate(line.text))
        elif (line.text.find('STATION') > -1):
            cell = assoc(cell, 'STATION', cleanStation(line.text))
        elif(line.text.find('MARKET') > -1):
            cell = assoc(cell, 'MARKET', cleanMarket(line.text))
        elif(cell.get('TITLE') != '' and cell.get('DATE') != ''
             and cell.get('STATION') != '' and cell.get('MARKET') != ''):
             break

    return cell

def groupInfo(filename):
    """Takes in a group document (multiple markets/stations) and returns a
       list of cells to be inserted"""
    cellList = []
    cell = {'TITLE': '', 'DATE': '', 'STATION': '', 'MARKET': ''}
    paras = deepcopy(Document(filename).paragraphs)
    for i in range(1,12):
        if (paras[i].text == ''):
            continue
        elif (cell.get('TITLE') == '' and paras[i].text != ''):
            cell = assoc(cell, 'TITLE', cleanTitle(paras[i].text))
        elif (paras[i].text.find('DATED') > -1):
            cell = assoc(cell, 'DATE', cleanDate(paras[i].text))

    #if Document(filename).tables is not None:
        #tbl = deepcopy(Document(filename).tables[0])

    docTable = deepcopy(Document(filename).tables[0])
    for i in range(len(docTable.rows)):
        mark = docTable.cell(i, 0).text[:docTable.cell(i, 0).text.find(':')]
        stat = docTable.cell(i, 0).text[docTable.cell(i, 0).text.find(':')+1:]
        cell = assoc(cell, 'MARKET', markDict.get(mark, mark))
        cell = assoc(cell, 'STATION', stat.strip())
        cellList.append(cell)
    for c in cellList:
        print(c)
    return cellList

#Stolen from Mary Rose Cook's "An introduction to functional programming"
#codewords.recurse.com/issues/one/an-introduction-to-functional-programming
def assoc(_d, key, value):
    from copy import deepcopy
    d = deepcopy(_d)
    d[key] = value
    return d

def cleanDate(strDate):
    """Converts date from written out format to mm/dd/yy"""
    strDate = strDate[strDate.find(':') + 2:].strip()

    mon = dateSwitch(strDate[0:strDate.find(' ')])

    dnum = strDate[strDate.find(' ')+1:strDate.find(',')]
    if len(dnum) == 1:
        dnum = '0' + dnum

    return mon + '/' + dnum + '/20'

def dateSwitch(strMon):
    """Convert written month to number date"""
    return monDict.get(strMon, 'ERROR')

def cleanStation(strStation):
    """Cut down para to just the station"""
    return strStation[strStation.find(':') + 1:strStation.find('AFFIL')].strip()

def cleanTitle(strTitle):
    """ Accounts for any title typos"""
    return titleDict.get(strTitle.strip(), strTitle)

def cleanMarket(strMarket):
    """ Accounts for any market typos to allow for proper insertion """
    strMarket = strMarket[strMarket.find(':')+2:strMarket.find('RA')-1].strip()
    return markDict.get(strMarket, strMarket)
